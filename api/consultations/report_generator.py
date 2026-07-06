import io
from xml.sax.saxutils import escape

from django.utils import timezone
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


def _display(value, fallback="Not recorded"):
    return str(value).strip() if value not in (None, "") else fallback


def _encounter_data(consultation):
    appointment = consultation.appointment
    patient = consultation.patient
    doctor = consultation.doctor
    diagnoses = list(consultation.diagnoses.all())
    prescriptions = list(consultation.prescriptions.prefetch_related("items").all())
    lab_requests = list(
        consultation.lab_requests.prefetch_related(
            "items__test_type",
            "items__result__verified_by",
        ).all()
    )

    return {
        "consultation": consultation,
        "appointment": appointment,
        "patient_name": _display(getattr(getattr(patient, "user", None), "full_name", None)),
        "patient_id": _display(getattr(patient, "patient_id", None)),
        "doctor_name": _display(getattr(getattr(doctor, "user", None), "full_name", None)),
        "diagnoses": diagnoses,
        "prescriptions": prescriptions,
        "lab_requests": lab_requests,
    }


def _add_docx_label(document, label, value):
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(_display(value))


def generate_encounter_docx(consultation):
    data = _encounter_data(consultation)
    appointment = data["appointment"]
    document = Document()

    title = document.add_heading("Clinical Encounter Record", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("DPAMS — Digital Patient Appointment Management System")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_docx_label(document, "Generated", timezone.localtime().strftime("%d %b %Y, %H:%M"))
    _add_docx_label(document, "Appointment ID", appointment.uuid)
    _add_docx_label(document, "Consultation ID", consultation.uuid)

    document.add_heading("Patient & Visit", level=1)
    _add_docx_label(document, "Patient", data["patient_name"])
    _add_docx_label(document, "Patient ID", data["patient_id"])
    _add_docx_label(document, "Attending doctor", data["doctor_name"])
    _add_docx_label(document, "Department / category", getattr(appointment.category, "name", None))
    _add_docx_label(document, "Scheduled date", appointment.appointment_date)
    _add_docx_label(document, "Scheduled time", appointment.start_time)
    _add_docx_label(document, "Consultation status", consultation.get_status_display())

    document.add_heading("Clinical Assessment", level=1)
    for label, value in (
        ("Chief complaint", consultation.chief_complaint),
        ("History of present illness", consultation.history_of_present_illness),
        ("Physical examination", consultation.physical_examination),
        ("Provisional diagnosis", consultation.provisional_diagnosis),
    ):
        _add_docx_label(document, label, value)

    document.add_heading("Diagnoses", level=1)
    if not data["diagnoses"]:
        document.add_paragraph("No diagnoses recorded.")
    for diagnosis in data["diagnoses"]:
        heading = diagnosis.disease_name
        if diagnosis.icd10_code:
            heading += f" ({diagnosis.icd10_code})"
        document.add_heading(heading, level=2)
        _add_docx_label(document, "Type", diagnosis.get_type_display())
        _add_docx_label(document, "Clinical notes", diagnosis.description)

    document.add_heading("Prescriptions", level=1)
    if not data["prescriptions"]:
        document.add_paragraph("No prescriptions issued.")
    for index, prescription in enumerate(data["prescriptions"], start=1):
        document.add_heading(f"Prescription {index}", level=2)
        _add_docx_label(document, "Notes", prescription.notes)
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["Medicine", "Dosage", "Frequency", "Duration", "Instructions"]
        for cell, text in zip(table.rows[0].cells, headers):
            cell.text = text
        for item in prescription.items.all():
            cells = table.add_row().cells
            values = [item.medicine_name, item.dosage, item.frequency, item.duration, item.instructions]
            for cell, value in zip(cells, values):
                cell.text = _display(value, "—")

    document.add_heading("Laboratory Requests & Results", level=1)
    if not data["lab_requests"]:
        document.add_paragraph("No laboratory tests requested.")
    for request in data["lab_requests"]:
        document.add_heading(f"Request {str(request.uuid)[:8]} — {request.get_status_display()}", level=2)
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for cell, text in zip(table.rows[0].cells, ["Test", "Result", "Remarks", "Verified by"]):
            cell.text = text
        for item in request.items.all():
            result = getattr(item, "result", None)
            cells = table.add_row().cells
            values = [
                item.test_type.name,
                getattr(result, "result", None) or "Pending",
                getattr(result, "remarks", None) or "—",
                getattr(getattr(result, "verified_by", None), "full_name", None) or "—",
            ]
            for cell, value in zip(cells, values):
                cell.text = _display(value, "—")

    document.add_paragraph()
    document.add_paragraph("Electronically generated clinical record.").italic = True
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def generate_encounter_pdf(consultation):
    data = _encounter_data(consultation)
    appointment = data["appointment"]
    buffer = io.BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title="Clinical Encounter Record",
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Muted", parent=styles["Normal"], textColor=colors.HexColor("#64748b"), fontSize=8))
    story = [
        Paragraph("Clinical Encounter Record", styles["Title"]),
        Paragraph("DPAMS — Digital Patient Appointment Management System", styles["Muted"]),
        Spacer(1, 8),
    ]

    def section(title):
        story.extend([Spacer(1, 8), Paragraph(title, styles["Heading2"]), Spacer(1, 4)])

    def details(rows):
        table = Table(
            [[Paragraph(f"<b>{escape(str(label))}</b>", styles["Normal"]), Paragraph(escape(_display(value)), styles["Normal"])] for label, value in rows],
            colWidths=[46 * mm, 128 * mm],
        )
        table.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.HexColor("#e2e8f0")),
        ]))
        story.append(table)

    details([
        ("Generated", timezone.localtime().strftime("%d %b %Y, %H:%M")),
        ("Appointment ID", appointment.uuid),
        ("Consultation ID", consultation.uuid),
    ])
    section("Patient & Visit")
    details([
        ("Patient", data["patient_name"]),
        ("Patient ID", data["patient_id"]),
        ("Attending doctor", data["doctor_name"]),
        ("Department / category", getattr(appointment.category, "name", None)),
        ("Scheduled date", appointment.appointment_date),
        ("Scheduled time", appointment.start_time),
        ("Status", consultation.get_status_display()),
    ])
    section("Clinical Assessment")
    details([
        ("Chief complaint", consultation.chief_complaint),
        ("History of present illness", consultation.history_of_present_illness),
        ("Physical examination", consultation.physical_examination),
        ("Provisional diagnosis", consultation.provisional_diagnosis),
    ])
    section("Diagnoses")
    if not data["diagnoses"]:
        story.append(Paragraph("No diagnoses recorded.", styles["Muted"]))
    else:
        diagnosis_rows = [["Diagnosis", "ICD-10", "Type", "Clinical notes"]]
        diagnosis_rows.extend([
            [item.disease_name, item.icd10_code or "—", item.get_type_display(), item.description or "—"]
            for item in data["diagnoses"]
        ])
        story.append(_pdf_table(diagnosis_rows, [45 * mm, 24 * mm, 27 * mm, 78 * mm], styles))

    section("Prescriptions")
    if not data["prescriptions"]:
        story.append(Paragraph("No prescriptions issued.", styles["Muted"]))
    for prescription in data["prescriptions"]:
        if prescription.notes:
            story.append(Paragraph(f"<b>Notes:</b> {escape(prescription.notes)}", styles["Normal"]))
        rows = [["Medicine", "Dosage", "Frequency", "Duration", "Instructions"]]
        rows.extend([
            [item.medicine_name, item.dosage, item.frequency, item.duration, item.instructions or "—"]
            for item in prescription.items.all()
        ])
        story.extend([_pdf_table(rows, [38 * mm, 27 * mm, 31 * mm, 27 * mm, 51 * mm], styles), Spacer(1, 8)])

    section("Laboratory Requests & Results")
    if not data["lab_requests"]:
        story.append(Paragraph("No laboratory tests requested.", styles["Muted"]))
    for request in data["lab_requests"]:
        story.append(Paragraph(f"<b>Request {str(request.uuid)[:8]}</b> — {request.get_status_display()}", styles["Normal"]))
        rows = [["Test", "Result", "Remarks", "Verified by"]]
        for item in request.items.all():
            result = getattr(item, "result", None)
            rows.append([
                item.test_type.name,
                getattr(result, "result", None) or "Pending",
                getattr(result, "remarks", None) or "—",
                getattr(getattr(result, "verified_by", None), "full_name", None) or "—",
            ])
        story.extend([_pdf_table(rows, [42 * mm, 42 * mm, 56 * mm, 34 * mm], styles), Spacer(1, 8)])

    story.extend([Spacer(1, 12), Paragraph("Electronically generated clinical record.", styles["Muted"])])
    document.build(story)
    buffer.seek(0)
    return buffer


def _pdf_table(rows, widths, styles):
    content = [
        [Paragraph(escape(_display(value, "—")), styles["Normal"]) for value in row]
        for row in rows
    ]
    table = Table(content, colWidths=widths, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f766e")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#cbd5e1")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("PADDING", (0, 0), (-1, -1), 5),
    ]))
    return table
