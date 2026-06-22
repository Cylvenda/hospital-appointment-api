import io
from django.utils import timezone
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

def generate_docx_report(lab_request):
    doc = Document()

    # Header
    heading = doc.add_heading('Laboratory Test Report', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Report Date: {timezone.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Request ID: {lab_request.uuid}")
    doc.add_paragraph("")

    # Patient and Doctor Details
    doc.add_heading('Patient Information', level=2)
    p = doc.add_paragraph()
    p.add_run('Patient Name: ').bold = True
    p.add_run(f"{lab_request.patient.user.full_name}\n")
    p.add_run('Requested By: ').bold = True
    p.add_run(f"Dr. {lab_request.doctor.user.full_name}\n")
    p.add_run('Date Requested: ').bold = True
    p.add_run(f"{lab_request.requested_at.strftime('%Y-%m-%d %H:%M')}")
    
    doc.add_paragraph("")

    # Results Table
    doc.add_heading('Test Results', level=2)
    
    items = lab_request.items.all()
    if not items:
        doc.add_paragraph("No tests were requested.")
    else:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header Row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Test Name'
        hdr_cells[1].text = 'Result'
        hdr_cells[2].text = 'Remarks'
        hdr_cells[3].text = 'Verified By'
        
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data Rows
        for item in items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.test_type.name
            
            result_obj = getattr(item, 'result', None)
            if result_obj:
                row_cells[1].text = result_obj.result
                row_cells[2].text = result_obj.remarks or ""
                row_cells[3].text = result_obj.verified_by.full_name if result_obj.verified_by else ""
            else:
                row_cells[1].text = "Pending"
                row_cells[2].text = "-"
                row_cells[3].text = "-"

    # Footer
    doc.add_paragraph("\n\n")
    doc.add_paragraph("This is an electronically generated report.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_pdf_report(lab_request):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    
    styles = getSampleStyleSheet()
    title_style = styles['Heading1']
    title_style.alignment = 1 # Center
    
    # Title
    elements.append(Paragraph("Laboratory Test Report", title_style))
    elements.append(Spacer(1, 12))
    
    # Details
    details_style = styles['Normal']
    elements.append(Paragraph(f"<b>Report Date:</b> {timezone.now().strftime('%Y-%m-%d %H:%M')}", details_style))
    elements.append(Paragraph(f"<b>Request ID:</b> {lab_request.uuid}", details_style))
    elements.append(Spacer(1, 12))
    
    elements.append(Paragraph("<b>Patient Information</b>", styles['Heading2']))
    elements.append(Paragraph(f"<b>Patient Name:</b> {lab_request.patient.user.full_name}", details_style))
    elements.append(Paragraph(f"<b>Requested By:</b> Dr. {lab_request.doctor.user.full_name}", details_style))
    elements.append(Paragraph(f"<b>Date Requested:</b> {lab_request.requested_at.strftime('%Y-%m-%d %H:%M')}", details_style))
    elements.append(Spacer(1, 12))
    
    # Table
    elements.append(Paragraph("<b>Test Results</b>", styles['Heading2']))
    elements.append(Spacer(1, 6))
    
    items = lab_request.items.all()
    if not items:
        elements.append(Paragraph("No tests were requested.", details_style))
    else:
        data = [['Test Name', 'Result', 'Remarks', 'Verified By']]
        
        for item in items:
            result_obj = getattr(item, 'result', None)
            if result_obj:
                data.append([
                    item.test_type.name,
                    result_obj.result,
                    result_obj.remarks or "-",
                    result_obj.verified_by.full_name if result_obj.verified_by else "-"
                ])
            else:
                data.append([item.test_type.name, "Pending", "-", "-"])
                
        table = Table(data, colWidths=[130, 100, 150, 100])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f3f4f6')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#1f2937')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb')),
        ]))
        elements.append(table)
        
    elements.append(Spacer(1, 30))
    elements.append(Paragraph("<i>This is an electronically generated report.</i>", details_style))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer
